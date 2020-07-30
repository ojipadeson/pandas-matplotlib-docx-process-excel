# -*- coding: utf-8 -*-
"""
Created on Wed Jul 29 11:07:55 2020

@author: MI
"""

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import datetime as dt
import matplotlib
matplotlib.use('Agg')

read_path = '数据示例_建筑业活动&工业品价格.xlsx'
out_path = '图片/'

plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

data_wind = pd.read_excel(read_path, sheet_name = 'data_wind', skiprows = [1, 2, 3, 4]).rename(columns = {'指标名称':'date'})
data_hand = pd.read_excel(read_path, sheet_name = 'data_手动', skiprows = 1)

#--------------------多年比较趋势图（共9张）------------------------
# 水泥磨机开工率---cement_operate_rate
cement_operate_rate = data_hand.iloc[:, 0:2].rename(columns = {'Unnamed: 0':'date'}).dropna()
# 水泥出货率---cement_shipment_rate
cement_shipment_rate = data_hand.iloc[:, 3:5].rename(columns = {'Unnamed: 3':'date'}).dropna()
# 水泥价格指数：全国---cement_price_index
cement_price_index = data_wind[['date', '水泥价格指数:全国']].dropna().reset_index(drop = True)
# 挖掘机工程量---engineer_quantity
engineer_quantity = data_hand.iloc[:, 10:12].rename(columns = {'Unnamed: 10':'date'}).dropna()
# 螺纹表观消费---apparent_consumption
apparent_consumption = data_hand.iloc[:, 13:15].rename(columns = {'Unnamed: 13':'date'}).dropna().reset_index(drop = True)
# 价格指数：普钢：螺纹---thread_price_index
thread_price_index = data_wind[['date', '价格指数:普钢:螺纹']].dropna().reset_index(drop = True)
# 全国建筑钢材成交量（周均）---steel_trans_volumn
steel_trans_volumn = data_hand.iloc[2:, [6, 8]].rename(columns = {'Unnamed: 6':'date'}).dropna().reset_index(drop = True)
# 线螺采购量（上海）---thread_buy
thread_buy = data_wind[['date', '线螺采购量:上海']].dropna().reset_index(drop = True)

def judge_trend(df):
    if 'year' in df:
        df = df.dropna().reset_index(drop = False).drop(columns = ['year']).set_index(['date'])
    else:
        df = df.dropna().reset_index(drop = False).set_index(['date'])
        
    if df.iloc[-1, 0] - df.iloc[-2, 0] > 0:
        return 1
    else:
        return 0

def multiple_lines(df, title, ylim = []):
    
    df['year'] = df.date.map(lambda x:x.year)
    
    for i in range(len(df['date'])):
        df.loc[i, 'date'] = dt.datetime.strptime('2020-' +\
                                                 str(df.loc[i, 'date'].month) +\
                                                 '-' + str(df.loc[i, 'date'].day),
                                                 '%Y-%m-%d')
    df = df.set_index(['year', 'date']).sort_index()
    
    fig, ax = plt.subplots(1, 1, figsize = (14, 7))
    for i in range(2015, 2021):
        df_year = df.query('year==%i' % i).iloc[::-1].reset_index(drop = False).drop(columns = ['year']).set_index(['date']).rename(columns = {title:'%i' % i})
        if i == 2020:
            plt.plot(df_year, linewidth = 5, color = 'black',label = '%i' % i)
        else:
            plt.plot(df_year, linewidth = 2.5, label = '%i' % i)
    plt.legend(loc = 'upper center', ncol = 6, fontsize = 15)
    
    xdate, xreal = [], []    
    for i in range(12):
        xdate.append(dt.datetime.strptime('2020-%i-1' % (i + 1), '%Y-%m-%d'))
        xreal.append('%i/1' % (i + 1))
    
    plt.yticks(size = 20)
    plt.xticks(xdate, xreal, size = 20)
    ax.set_xlim(dt.datetime.strptime('2020-1-1', '%Y-%m-%d'), None)
    #plt.title(title, fontsize = 'xx-large', fontweight = 'bold')
    
    if ylim != []:
        ax.set_ylim(ylim[0], ylim[1])
    
    if judge_trend(df.query('year==2020')) == 1:
        plt.savefig(out_path + title + '上升.png')
    else:
        plt.savefig(out_path + title + '下降.png')

# 请注意excel表格中日期格式
# 删除冒号

multiple_lines(cement_operate_rate, '水泥磨机开工率')
multiple_lines(cement_shipment_rate, '水泥出货率')
multiple_lines(engineer_quantity, '挖掘机工程量', ylim = [0, 20000])
multiple_lines(steel_trans_volumn, '全国建筑钢材成交量（周均）')
multiple_lines(cement_price_index, '水泥价格指数全国')
multiple_lines(thread_price_index, '价格指数普钢螺纹')
multiple_lines(thread_buy, '线螺采购量上海')
# 这一幅图是重复的，重新读取原始数据
cement_price_index = data_wind[['date', '水泥价格指数:全国']].dropna().reset_index(drop = True)
multiple_lines(cement_price_index, '水泥价格全国')
multiple_lines(apparent_consumption, '螺纹表观消费')


#--------------------多年单线趋势图（共8张）------------------------
# CRB现货指数：综合---com_crb
com_crb = data_wind[['date', 'CRB现货指数:综合']].dropna().reset_index(drop = True)
# 生产资料价格指数---product_price_index
product_price_index = data_wind[['date', '生产资料价格指数']].dropna().reset_index(drop = True)
# 中国大宗商品价格指数：总指数---commodity_price_index
commodity_price_index = data_wind[['date', '中国大宗商品价格指数:总指数']].dropna().reset_index(drop = True)
# 现货价原油：英国布伦特---british_brent
british_brent = data_wind[['date', '现货价:原油:英国布伦特Dtd']].dropna().reset_index(drop = True)
# 中国大宗商品价格指数:矿产类---mineral_index
mineral_index = data_wind[['date', '中国大宗商品价格指数:矿产类']].dropna().reset_index(drop = True)
# 中国大宗商品价格指数:钢铁类---steel_index
steel_index = data_wind[['date', '中国大宗商品价格指数:钢铁类']].dropna().reset_index(drop = True)
# 中国大宗商品价格指数:有色类---colored_index
colored_index = data_wind[['date', '中国大宗商品价格指数:有色类']].dropna().reset_index(drop = True)
# 中国化工产品价格指数---chemical_index
chemical_index = data_wind[['date', '中国化工产品价格指数(CCPI)']].dropna().reset_index(drop = True)

def one_line(df, title, xlim = [], ylim = []):
    df = df.set_index(['date'])
    df = df.sort_index()
    if xlim == []:
        df = df['2019-06-03':df.index[-1]]
        xlim = ['2019-06-03', df.index[-1]]
    else:
        df = df[xlim[0]:df.index[-1]]
    
    fig, ax = plt.subplots(1, 1, figsize = (14, 7))
    plt.plot(df, linewidth = 2.5, label = title)
    
    plt.legend(loc = 'upper center', fontsize = 15)
    
    if ylim != []:
        ax.set_ylim(ylim[0], ylim[1])
    
    ax.set_xlim(dt.datetime.strptime(xlim[0], '%Y-%m-%d'), None)
    #plt.title(title, fontsize = 'x-large', fontweight = 'bold')
    plt.yticks(size = 20)
    plt.xticks(size = 20)
    
    if judge_trend(df) == 1:
        plt.savefig(out_path + title + '上升.png')
    else:
        plt.savefig(out_path + title + '下降.png')

one_line(com_crb, 'CRB现货指数综合', xlim = ['2018-12-03'])
one_line(product_price_index, '生产资料价格指数')
one_line(commodity_price_index, '中国大宗商品价格指数总指数')
one_line(british_brent, '现货价原油英国布伦特Dtd', xlim = ['2018-12-03'])
one_line(mineral_index,'中国大宗商品价格指数矿产类', ylim = [100, 170])
one_line(steel_index, '中国大宗商品价格指数钢铁类', ylim = [90, 120])
one_line(colored_index, '中国大宗商品价格指数有色类', ylim = [60, 85])
one_line(chemical_index, '中国化工产品价格指数(CCPI)', xlim = ['2018-12-03'])


#--------------------双轴比较趋势图（共7张）------------------------
# 计算同比和环比
def calculate_th(df):
    avg = df.set_index(['date']).resample('M').mean()
    t = avg.groupby(avg.index.map(lambda x:x.month)).apply(lambda x:x.pct_change()) * 100
    h = df.set_index(['date']).resample('M').mean().pct_change() * 100
    return t, h

# df1为同比，df2为环比
def double_y(df1, df2, title, ylim1, ylim2, flag = 0, begin_date = '2015-08'):
    df1 = df1.sort_index()
    df1 = df1[begin_date:df1.index[-1]]
    df2 = df2.sort_index()
    df2 = df2[begin_date:df2.index[-1]]
    
    fig = plt.figure(figsize = (14, 7))
    ax = fig.add_subplot(111)
    
    label = []
    if flag == 1:
        label = ['现货', '期货']
    elif flag == 0:
        label = ['同比', '环比']
    
    lns1 = ax.plot(df1, '-', linewidth = 2.5, label = label[0])
    ax2 = ax.twinx()
    lns2 = ax2.plot(df2, '-r', linewidth = 2.5, label = label[1])
    lns = lns1+lns2
    labs = [l.get_label() for l in lns]
    ax.legend(lns, labs, loc='upper center', ncol = 2, fontsize = 15)
    if flag == 0:
        ax.spines['top'].set_color('none')
        ax.xaxis.set_ticks_position('bottom')
        ax.spines['bottom'].set_position(('data', 0))
    ax.set_ylim(ylim1[0], ylim1[1])
    ax2.set_ylim(ylim2[0], ylim2[1])
        
    #plt.title(title + label[0] + label[1], fontsize = 'x-large', fontweight = 'bold')
    
    if judge_trend(df1) == 1 and judge_trend(df2) == 1:
        plt.savefig(out_path + title + label[0] + '上升、' + label[1] + '上升.png')
    elif judge_trend(df1) == 1 and judge_trend(df2) == 0:
        plt.savefig(out_path + title + label[0] + '上升、' + label[1] + '下降.png')
    elif judge_trend(df1) == 0 and judge_trend(df2) == 1:
        plt.savefig(out_path + title + label[0] + '下降、' + label[1] + '上升.png')
    else:
        plt.savefig(out_path + title + label[0] + '下降、' + label[1] + '下降.png')

# 同比环比共3张

yoy, chain_ratio = calculate_th(com_crb)
double_y(yoy, chain_ratio, 'CRB现货指数', [-20, 20], [-10, 5]) # ylim1为左轴
yoy, chain_ratio = calculate_th(product_price_index)
double_y(yoy, chain_ratio, '生产资料价格指数', [-20, 30], [-6, 6])
yoy, chain_ratio = calculate_th(commodity_price_index)
double_y(yoy, chain_ratio, '大宗商品价格指数（总指数）', [-40, 80], [-15, 20])

# 现货期货共4张
# 铜现货---spot_copper，南华沪铜指数---future_copper
spot_copper = data_wind[['date', '现货价:铜:1#:全国']].dropna().reset_index(drop = True).set_index(['date'])
future_copper = data_wind[['date', '南华沪铜指数']].dropna().reset_index(drop = True).set_index(['date'])
# 塑料现货---spot_plastic，南华塑料指数---future_plastic
spot_plastic = data_wind[['date', '中国塑料现货价格指数(中塑现货指数)']].dropna().reset_index(drop = True).set_index(['date'])
future_plastic = data_wind[['date', '南华塑料指数']].dropna().reset_index(drop = True).set_index(['date'])
# 动力煤现货---spot_coal，南华动力煤指数---future_coal
spot_coal = data_wind[['date', '平均价:动力煤:国内主要地区']].dropna().reset_index(drop = True).set_index(['date'])
future_coal = data_wind[['date', '南华动力煤指数']].dropna().reset_index(drop = True).set_index(['date'])
# 玻璃现货---spot_glass，南华玻璃指数---future_glass
spot_glass = data_wind[['date', '中国玻璃价格指数']].dropna().reset_index(drop = True).set_index(['date'])
future_glass = data_wind[['date', '南华玻璃指数']].dropna().reset_index(drop = True).set_index(['date'])

double_y(spot_copper, future_copper, '铜', [37000, 53000], [2500, 4500], flag = 1, begin_date = '2019-06')
double_y(spot_plastic, future_plastic, '塑料', [850, 1050], [600, 1000], flag = 1, begin_date = '2019-06')
double_y(spot_coal, future_coal, '动力煤', [450, 610], [1300, 1650], flag = 1, begin_date = '2019-06')
double_y(spot_glass, future_glass, '玻璃', [900, 1250], [1300, 2100], flag = 1, begin_date = '2019-06')

#--------------------转移图片至指定文档------------------------
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import os

document = Document()

sections = document.sections
for section in sections:
    section.left_margin = 0
    section.right_margin = 0

table_rows = (8 + 16 + 3)
table_cols = 2

table1 = document.add_table(rows=table_rows, cols=table_cols)

folder_path = '图片/'

name_list = ['水泥磨机开工率', '水泥出货率',
              '水泥价格指数全国', '挖掘机工程量',
              '螺纹表观消费', '价格指数普钢螺纹',
              '全国建筑钢材成交量（周均）', '线螺采购量上海',
              'CRB现货指数综合', 'CRB现货指数',
              '生产资料价格指数', '生产资料价格指数',
              '中国大宗商品价格指数总指数', '大宗商品价格指数（总指数）',
              '现货价原油英国布伦特Dtd', '中国大宗商品价格指数矿产类',
              '中国大宗商品价格指数钢铁类', '中国大宗商品价格指数有色类',
              '中国化工产品价格指数(CCPI)', '水泥价格全国',
              '铜', '塑料',
              '动力煤', '玻璃']

quota_list = []

def extract(string, sort = 1):
    if sort == 1:
        if os.path.exists(folder_path + string + '上升.png') == True:
            string = string + '上升'
        else:
            string = string + '下降'
    elif sort == 2:
        if os.path.exists(folder_path + string + '同比上升、环比上升.png') == True:
            string = string + '同比上升、环比上升'
        elif os.path.exists(folder_path + string + '同比上升、环比下降.png') == True:
            string = string + '同比上升、环比下降'
        elif os.path.exists(folder_path + string + '同比下降、环比下降.png') == True:
            string = string + '同比下降、环比下降'
        else:
            string = string + '同比下降、环比上升'
    else:
        if os.path.exists(folder_path + string + '现货上升、期货上升.png') == True:
            string = string + '现货上升、期货上升'
        elif os.path.exists(folder_path + string + '现货上升、期货下降.png') == True:
            string = string + '现货上升、期货下降'
        elif os.path.exists(folder_path + string + '现货下降、期货下降.png') == True:
            string = string + '现货下降、期货下降'
        else:
            string = string + '现货下降、期货上升'
    quota_list.append(string)
    
for i in range(24):
    if i == 9 or i == 11 or i == 13:
        extract(name_list[i], sort = 2)
    elif i >= 20:
        extract(name_list[i], sort = 3)
    else:
        extract(name_list[i])

def set_table_picture(row, line, read_path, height = 2.3):
    run = document.tables[0].cell(row, line).paragraphs[0].add_run()
    run.add_picture(read_path, width = Inches(4.4), height = Inches(height))
    
def set_table_paragraph(row, line, paragraph):
    run = table1.cell(row, line).paragraphs[0].add_run(paragraph)
    run.font.name='宋体' 
    run.font.size=Pt(10)
    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.bold=1         
    run.font.color.rgb=RGBColor.from_string('000000')
    table1.cell(row, line).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

run = table1.cell(0, 0).paragraphs[0].add_run(' ' * 5 + '一、建筑业活动')
run.font.name='宋体' 
run.font.size=Pt(18)
run.font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.bold=1         
run.font.color.rgb=RGBColor.from_string('000000')
#table1.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

run = table1.cell(9, 0).paragraphs[0].add_run(' ' * 5 + '二、工业品价格')
run.font.name='宋体' 
run.font.size=Pt(18)
run.font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.bold=1         
run.font.color.rgb=RGBColor.from_string('000000')
#table1.cell(9, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for i in range(1, 5):
    set_table_paragraph(2 * i - 1, 0, quota_list[i * 2 - 2])
    set_table_picture(2 * i, 0, folder_path + quota_list[2 * i - 2] + '.png')
    set_table_paragraph(2 * i - 1, 1, quota_list[i * 2 - 1])
    set_table_picture(2 * i, 1, folder_path + quota_list[2 * i - 1] + '.png')
    
for i in range(5, 13):
    if i * 2 >= 20:
        set_table_paragraph(2 * i + 1, 0, quota_list[i * 2 - 2])
        set_table_picture(2 * i + 2, 0, folder_path + quota_list[2 * i - 2] + '.png', 2.31)
        set_table_paragraph(2 * i + 1, 1, quota_list[i * 2 - 1])
        set_table_picture(2 * i + 2, 1, folder_path + quota_list[2 * i - 1] + '.png', 2.31)
    else:
        set_table_paragraph(2 * i, 0, quota_list[i * 2 - 2])
        set_table_picture(2 * i + 1, 0, folder_path + quota_list[2 * i - 2] + '.png', 2.31)
        set_table_paragraph(2 * i, 1, quota_list[i * 2 - 1])
        set_table_picture(2 * i + 1, 1, folder_path + quota_list[2 * i - 1] + '.png', 2.31)

#document.add_page_break()

document.save('test.docx')